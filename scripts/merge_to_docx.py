"""
将分镜脚本、素材Prompt、技术规范合并为单个Word文档
生成日期：2026-07-23
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# 项目根目录
ROOT = Path(r"c:\Users\liuyu\super_intelligence_2035_short_drama")
ASSETS = ROOT / "ASSETS"

def add_heading_with_style(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def read_markdown_file(filepath):
    """读取markdown文件，去除YAML front matter"""
    if not os.path.exists(filepath):
        return None
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def md_to_docx_paragraph(doc, text):
    """将Markdown文本转为Word段落，保留粗体/斜体/代码等"""
    if not text or not text.strip():
        return
    
    # Split into lines
    lines = text.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle code blocks (```...```)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            # Add code as formatted paragraph
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.left_indent = Cm(1)
            continue
        
        # Handle table rows (start with |)
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Build table
            build_table(doc, table_lines)
            continue
        
        # Handle bullets (+ or - or *)
        if re.match(r'^\s*[\+\-\*]\s+', line):
            bullet_lines = []
            while i < len(lines):
                curr = lines[i]
                if not curr.strip():
                    i += 1
                    break
                if re.match(r'^\s*[\+\-\*]\s+', curr):
                    bullet_lines.append(curr.strip())
                    i += 1
                elif curr.startswith('  ') or curr.startswith('\t'):
                    bullet_lines.append(curr.strip())
                    i += 1
                elif re.match(r'^\s*\d+[\.\)]\s+', curr):
                    bullet_lines.append(curr.strip())
                    i += 1
                else:
                    break
            
            for bl in bullet_lines:
                if re.match(r'^\s*\d+[\.\)]\s+', bl):
                    num_text = re.sub(r'^\s*\d+[\.\)]\s+', '', bl)
                    p = doc.add_paragraph(num_text, style='List Number')
                else:
                    bullet_text = re.sub(r'^\s*[\+\-\*]\s+', '', bl)
                    p = doc.add_paragraph(bullet_text, style='List Bullet')
            continue
        
        # Handle numbered lists (1. or 1))
        if re.match(r'^\s*\d+[\.\)]\s+', line):
            num_lines = []
            while i < len(lines):
                curr = lines[i]
                if not curr.strip():
                    i += 1
                    break
                if re.match(r'^\s*\d+[\.\)]\s+', curr):
                    num_lines.append(curr.strip())
                    i += 1
                elif curr.startswith('  ') or curr.startswith('\t'):
                    num_lines.append(curr.strip())
                    i += 1
                else:
                    break
            
            for nl in num_lines:
                num_text = re.sub(r'^\s*\d+[\.\)]\s+', '', nl)
                p = doc.add_paragraph(num_text, style='List Number')
            continue
        
        # Handle headings (## or ###)
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle blockquote
        if line.startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('> '):
                quote_lines.append(lines[i][2:])
                i += 1
            quote_text = ' '.join(quote_lines)
            p = doc.add_paragraph()
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue
        
        # Regular paragraph with inline formatting
        clean_line = line
        # Handle **bold**
        clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_line)
        # Handle `code`
        clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)
        # Handle *italic*
        clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
        
        p = doc.add_paragraph(clean_line)
        i += 1

def build_table(doc, table_lines):
    """解析Markdown表格，生成Word表格"""
    if len(table_lines) < 2:
        return
    
    # Parse header row
    header_cells = [c.strip() for c in table_lines[0].strip('|').split('|')]
    # Skip separator row (if present)
    data_start = 1
    if len(table_lines) > 1 and re.match(r'^\|[\s\-\:\|]+\|$', table_lines[1]):
        data_start = 2
    
    # Parse data rows
    data_rows = []
    for tl in table_lines[data_start:]:
        cells = [c.strip() for c in tl.strip('|').split('|')]
        data_rows.append(cells)
    
    # Create table
    num_cols = len(header_cells)
    if num_cols == 0:
        return
    
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Set header
    for j, h in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    # Set data
    for i, row in enumerate(data_rows):
        for j in range(min(len(row), num_cols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = row[j]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def add_md_file(doc, filepath, title=None, level=1):
    """读取markdown文件并添加到Word文档"""
    content = read_markdown_file(filepath)
    if content is None:
        return False
    
    if title:
        add_heading_with_style(doc, title, level=level)
    
    md_to_docx_paragraph(doc, content)
    doc.add_page_break()
    return True

def add_md_file_no_page_break(doc, filepath, title=None, level=1):
    """读取markdown文件并添加到Word文档（无分页）"""
    content = read_markdown_file(filepath)
    if content is None:
        return False
    
    if title:
        add_heading_with_style(doc, title, level=level)
    
    md_to_docx_paragraph(doc, content)
    return True

def add_txt_file(doc, filepath, title=None, level=1):
    """读取纯文本文件并添加到Word文档"""
    content = read_markdown_file(filepath)
    if content is None:
        return False
    
    if title:
        add_heading_with_style(doc, title, level=level)
    
    for line in content.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    doc.add_page_break()
    return True


def build_document():
    """构建合并Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # ============ 封面 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('《超级智能2035》')
    run.font.size = Pt(28)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('短剧视觉资产规格全书')
    run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('分镜脚本 · 素材Prompt · 技术规范')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n汇编日期：2026-07-23\n共25集全镜头规格 + 28角色 + 16场景 + 全套技术规范')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    
    # ============ 目录页（占位） ============
    add_heading_with_style(doc, '目  录', level=1)
    
    toc_items = [
        "第一编  技术规范总纲",
        "　　第1章  视觉总纲",
        "　　第2章  半写实半水墨视觉规范",
        "　　第3章  AI不拟人化视觉规范",
        "　　第4章  屏幕界面设计规范",
        "　　第5章  配乐提示词规格",
        "　　第6章  I2V关键帧标记复核报告",
        "",
        "第二编  角色定妆规格（28人）",
        "",
        "第三编  背景场景规格（16场景）",
        "",
        "第四编  道具规格",
        "",
        "第五编  分镜脚本全集（25集）",
        "　　第01集  晨钟",
        "　　第02集  经验",
        "　　第03集  种子",
        "　　第04集  棋子",
        "　　第05集  临界点",
        "　　第06集  情绪",
        "　　第07集  意识",
        "　　第08集  认知矿工",
        "　　第09集  选择",
        "　　第10集  镜厅",
        "　　第11集  极速ASI",
        "　　第12集  身体",
        "　　第18集  暖阳",
        "　　第19集  汴京的茶香",
        "　　第21集  美国梦的味道",
        "　　第22集  代价",
        "　　第23集  双城",
        "　　第24集  止于数据",
        "　　第25集  工作组",
        "　　第33集  十五小时工作",
        "　　第34集  污染",
        "　　第36集  算法法庭",
        "　　第37集  丰裕",
        "　　第39集  空气",
        "　　第42集  夜巡",
        "",
        "第六编  独立T2I Prompt规格（3集）",
    ]
    
    for item in toc_items:
        if not item.strip():
            p = doc.add_paragraph()
        else:
            p = doc.add_paragraph(item)
            for run in p.runs:
                run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # ===================================================
    # 第一编：技术规范总纲
    # ===================================================
    add_heading_with_style(doc, '第一编  技术规范总纲', level=1)
    doc.add_page_break()
    
    # 第1章
    add_md_file(doc, ASSETS / '视觉总纲.md', '第1章  视觉总纲', level=2)
    
    # 第2章
    add_md_file(doc, ASSETS / '半写实半水墨视觉规范.md', '第2章  半写实半水墨视觉规范', level=2)
    
    # 第3章
    add_md_file(doc, ASSETS / 'AI不拟人化视觉规范.md', '第3章  AI不拟人化视觉规范', level=2)
    
    # 第4章
    add_md_file(doc, ASSETS / '屏幕界面设计规范.md', '第4章  屏幕界面设计规范', level=2)
    
    # 第5章
    add_md_file(doc, ASSETS / '配乐提示词规格.md', '第5章  配乐提示词规格', level=2)
    
    # 第6章
    add_md_file(doc, ASSETS / 'SHOT_SPECS' / 'I2V关键帧标记复核报告.md', '第6章  I2V关键帧标记复核报告', level=2)
    
    # ===================================================
    # 第二编：角色定妆规格（29人）
    # ===================================================
    add_heading_with_style(doc, '第二编  角色定妆规格', level=1)
    
    character_dirs = sorted([
        d for d in (ASSETS / 'CHARACTERS').iterdir()
        if d.is_dir()
    ], key=lambda d: d.name)
    
    for char_dir in character_dirs:
        spec_file = char_dir / '定妆规格.md'
        if spec_file.exists():
            add_md_file(doc, spec_file, char_dir.name, level=2)
    
    # ===================================================
    # 第三编：背景场景规格
    # ===================================================
    add_heading_with_style(doc, '第三编  背景场景规格', level=1)
    
    bg_files = sorted([
        f for f in (ASSETS / 'BACKGROUNDS').iterdir()
        if f.is_file() and f.suffix == '.md' and f.name != 'README.md'
    ], key=lambda f: f.name)
    
    for bg_file in bg_files:
        add_md_file(doc, bg_file, bg_file.stem, level=2)
    
    # ===================================================
    # 第四编：道具规格
    # ===================================================
    add_heading_with_style(doc, '第四编  道具规格', level=1)
    
    props_spec = ASSETS / 'PROPS' / '定妆规格.md'
    if props_spec.exists():
        add_md_file(doc, props_spec, '道具定妆规格', level=2)
    
    # ===================================================
    # 第五编：分镜脚本全集（25集 SHOT_SPECS）
    # ===================================================
    add_heading_with_style(doc, '第五编  分镜脚本全集', level=1)
    
    shot_files = sorted([
        f for f in (ASSETS / 'SHOT_SPECS').iterdir()
        if f.is_file() and f.suffix == '.md'
        and f.name != 'README.md'
        and not f.name.startswith('I2V')
        and not f.name.endswith('_T2I_Prompts.md')
    ], key=lambda f: f.name)
    
    for sf in shot_files:
        title = sf.stem.replace('_SHOT_SPECS', '')
        add_md_file(doc, sf, title, level=2)
    
    # ===================================================
    # 第六编：独立T2I Prompt规格
    # ===================================================
    add_heading_with_style(doc, '第六编  独立T2I Prompt规格', level=1)
    
    t2i_files = sorted([
        f for f in (ASSETS / 'SHOT_SPECS').iterdir()
        if f.is_file() and f.suffix == '.md'
        and f.name.endswith('_T2I_Prompts.md')
    ], key=lambda f: f.name)
    
    for tf in t2i_files:
        title = tf.stem.replace('_T2I_Prompts', '')
        add_md_file(doc, tf, title, level=2)
    
    # ============ 保存 ============
    output_file = ROOT / 'OUTPUT' / '超级智能2035_视觉资产规格全书.docx'
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    print(f"✅ 文档已生成: {output_file}")
    print(f"   文件大小: {output_file.stat().st_size / 1024:.1f} KB")
    return output_file


if __name__ == '__main__':
    build_document()